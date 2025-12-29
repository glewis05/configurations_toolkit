"""
Word Document Parser for Clinic Specification Documents

PURPOSE: Parse clinic specification Word documents (.docx) to extract
         configuration settings, providers, and locations

R EQUIVALENT: Like readxl::read_xlsx but for Word docs with table extraction

AVIATION ANALOGY: Like reading a mission briefing document -
extract the key parameters from a standardized format

EXPECTED DOCUMENT STRUCTURE:
- Header with Doc ID, Version, Parent SRS
- Purpose and Scope sections
- Configuration Matrix table with columns:
  - Category
  - Global Default
  - [Clinic Name] Override (or Customization)
  - Rationale / Source
- Provider information (may be in table or structured text)
- Change Log table

AUTHOR: Glen Lewis
DATE: 2024
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime

# python-docx is the main library for reading Word documents
# Install with: pip install python-docx
try:
    from docx import Document
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Run: pip install python-docx")


# =============================================================================
# CATEGORY TO CONFIG KEY MAPPING
# =============================================================================
# This mapping translates document category names to database config keys.
# The document uses human-readable category names; the database uses normalized keys.
#
# Each category can map to multiple config keys because some document rows
# contain multiple pieces of data that need to be split apart.
#
# NOTE: Location-specific values (helpdesk_phone, signature_block, hours)
# are parsed with 'parse_location_specific' which distributes values to
# actual location records using fuzzy name matching.
# =============================================================================

CATEGORY_TO_CONFIG_KEYS = {
    # Patient Appointment Extract – Filtering
    # This category contains multiple values that need to be parsed separately:
    # - Patient Status: New
    # - Providers: ... (by location)
    # NOTE: extract_service_locations and extract_appointment_types removed -
    # these were not useful config keys
    'patient appointment extract': {
        'keys': ['extract_patient_status', 'extract_providers'],
        'parser': 'parse_appointment_extract'  # Special parser method
    },

    # Invitation Schedule (Text & Email)
    'invitation schedule': {
        'keys': ['invitation_days_before', 'invitation_channels'],
        'parser': 'simple'
    },

    # NOTE: email_branding category removed - not used in production

    # Invitation and SMS Signature Line (Email/SMS)
    # These are location-specific values
    'signature': {
        'keys': ['signature_block_email', 'signature_block_sms'],
        'parser': 'parse_location_specific'  # Has per-location values
    },

    # Assessment Lockout
    'assessment lockout': {
        'keys': ['assessment_lockout_trigger', 'assessment_lockout_period_months'],
        'parser': 'simple'
    },

    # Invitation Clinical Help Desk Email
    'help desk email': {
        'keys': ['helpdesk_email', 'helpdesk_workflow'],
        'parser': 'parse_helpdesk_email'  # Special: separate email from workflow
    },
    'helpdesk email': {
        'keys': ['helpdesk_email', 'helpdesk_workflow'],
        'parser': 'parse_helpdesk_email'
    },

    # Invitation Clinical Help Desk Phone
    # Location-specific: each location has its own phone number
    'help desk phone': {
        'keys': ['helpdesk_phone'],
        'parser': 'parse_location_specific'  # May have per-location phones
    },
    'helpdesk phone': {
        'keys': ['helpdesk_phone'],
        'parser': 'parse_location_specific'
    },

    # Invitation Clinical Hours of Operations
    # Location-specific: each location has its own hours
    'hours of operation': {
        'keys': ['hours_open', 'hours_close'],
        'parser': 'parse_location_specific_hours'
    },
    'operating hours': {
        'keys': ['hours_open', 'hours_close'],
        'parser': 'parse_location_specific_hours'
    },

    # TC Module
    'tc module': {
        'keys': ['tc_scoring_enabled'],
        'parser': 'simple'
    },

    # TC Age Range
    'tc age': {
        'keys': ['tc_minimum_age', 'tc_maximum_age'],
        'parser': 'parse_age_range'
    },
    'age range': {
        'keys': ['tc_minimum_age', 'tc_maximum_age'],
        'parser': 'parse_age_range'
    },

    # Default Lab Order
    'default lab order': {
        'keys': ['lab_default_test_code', 'lab_default_test_name'],
        'parser': 'parse_lab_order'
    },
    'lab order': {
        'keys': ['lab_default_test_code', 'lab_default_test_name'],
        'parser': 'parse_lab_order'
    },

    # Default Lab
    'default lab': {
        'keys': ['lab_default_name'],
        'parser': 'simple'
    },

    # Default Sample
    'default sample': {
        'keys': ['lab_default_sample'],
        'parser': 'simple'
    },

    # Optional Sample Options
    'optional sample': {
        'keys': ['lab_optional_samples'],
        'parser': 'simple'
    },

    # Optional Lab Tests
    'optional lab': {
        'keys': ['lab_optional_tests'],
        'parser': 'simple'
    },

    # Default Provider Information - goes to providers table, not config
    # Now maps to extract_providers config key for location-specific providers
    'provider information': {
        'keys': ['extract_providers'],
        'parser': 'providers'
    },
    'default provider': {
        'keys': ['extract_providers'],
        'parser': 'providers'
    },
}


class ClinicSpecParser:
    """
    PURPOSE: Parse clinic specification Word documents

    Expected document structure:
    - Header with Doc ID, Version, Parent SRS
    - Purpose and Scope sections
    - Configuration Matrix table with columns:
      - Category
      - Global Default
      - [Clinic Name] Override (or Customization)
      - Rationale / Source
    - Change Log table

    AVIATION ANALOGY: Like reading a mission briefing document -
    extract the key parameters from a standardized format.

    PARAMETERS:
        file_path: Path to the Word document (.docx)

    EXAMPLE:
        parser = ClinicSpecParser("Portland_Clinic_Spec.docx")
        result = parser.parse()
        print(result['clinic_name'])  # 'Portland'
        print(result['configurations'][0])  # First config row
    """

    def __init__(self, file_path: str):
        """
        Load Word document.

        WHY THIS APPROACH: We load the document once and parse it
        in sections. This is more efficient than re-reading for
        each parse operation.
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        self.file_path = Path(file_path)

        if not self.file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        if not self.file_path.suffix.lower() == '.docx':
            raise ValueError(f"Expected .docx file, got: {self.file_path.suffix}")

        # Load the document
        self.doc = Document(str(self.file_path))

        # Initialize result containers
        self.result = {
            'doc_id': None,
            'version': None,
            'parent_srs': None,
            'clinic_name': None,
            'program': None,
            'scope_locations': [],
            'configurations': [],      # Raw parsed configs
            'mapped_configs': [],      # Configs mapped to config_keys
            'providers': [],
            'change_log': [],
            'raw_tables': []  # Store raw table data for debugging
        }

    def parse(self) -> Dict:
        """
        Parse the entire document.

        PURPOSE: Main entry point - extracts all structured data

        RETURNS:
            Dict containing:
            - doc_id: Document identifier (e.g., 'P4M-CL-PORT-SPEC')
            - version: Document version (e.g., '0.1')
            - parent_srs: Parent SRS reference (e.g., 'P4M-SRS_v3.0')
            - clinic_name: Extracted clinic name (e.g., 'Portland')
            - program: Program name (e.g., 'Prevention4ME')
            - scope_locations: List of location names from Scope section
            - configurations: List of raw config dicts from config table
            - mapped_configs: List of configs mapped to config_keys
            - providers: List of provider dicts
            - change_log: List of change log entries

        WHY THIS APPROACH: We parse in a specific order because
        some sections help interpret others (e.g., doc_id helps
        identify the clinic name).
        """
        # Parse header info first (doc_id, version, etc.)
        self._parse_header()

        # Extract clinic name from doc_id or content
        self._extract_clinic_name()

        # Parse scope section for location names
        self._parse_scope()

        # Parse all tables
        self._parse_all_tables()

        # Map raw configs to config_keys
        self._map_configs_to_keys()

        return self.result

    # =========================================================================
    # CONFIG KEY MAPPING - Core new functionality
    # =========================================================================

    def _map_configs_to_keys(self) -> None:
        """
        Map raw configurations to specific config_keys.

        PURPOSE: Transform document category names to database config keys
                 and parse complex cell values into separate configs.

        WHY THIS APPROACH: The document uses human-readable category names
        but the database uses normalized keys. Some categories contain
        multiple values that need to be split into separate config records.
        """
        for config in self.result['configurations']:
            category = config.get('category', '')
            override = config.get('override', '')
            global_default = config.get('global_default', '')
            rationale = config.get('rationale', '')

            # Find matching category mapping
            mapping = self._find_category_mapping(category)

            if not mapping:
                # No mapping found - store raw for manual review
                self.result['mapped_configs'].append({
                    'config_key': f'unmapped_{self._slugify(category)}',
                    'value': override or global_default,
                    'source_category': category,
                    'rationale': rationale,
                    'is_unmapped': True
                })
                continue

            parser_type = mapping.get('parser', 'simple')
            config_keys = mapping.get('keys', [])

            # Skip provider rows (handled separately)
            if parser_type == 'providers':
                continue

            # Parse based on parser type
            # NOTE: Location-specific parsers distribute values to actual location
            # records using fuzzy name matching against self.result['scope_locations']
            if parser_type == 'parse_appointment_extract':
                parsed = self._parse_appointment_extract(override, global_default)
            elif parser_type == 'parse_helpdesk_email':
                parsed = self._parse_helpdesk_email(override, global_default)
            elif parser_type == 'parse_location_specific':
                parsed = self._parse_location_specific_value(override, global_default, config_keys)
            elif parser_type == 'parse_location_specific_hours':
                parsed = self._parse_location_specific_hours(override, global_default, config_keys)
            elif parser_type == 'parse_hours':
                parsed = self._parse_hours(override, global_default)
            elif parser_type == 'parse_age_range':
                parsed = self._parse_age_range(override, global_default)
            elif parser_type == 'parse_lab_order':
                parsed = self._parse_lab_order(override, global_default)
            else:
                # Simple: use first config key, store whole value
                parsed = {config_keys[0]: override or global_default} if config_keys else {}

            # Add parsed configs to mapped_configs
            for config_key, value in parsed.items():
                if value:  # Only add non-empty values
                    self.result['mapped_configs'].append({
                        'config_key': config_key,
                        'value': value,
                        'source_category': category,
                        'rationale': rationale,
                        'is_unmapped': False
                    })

    def _find_category_mapping(self, category: str) -> Optional[Dict]:
        """
        Find the mapping for a given category.

        PURPOSE: Match document category to our mapping table.

        Uses partial matching because document categories may have
        extra text like "Patient Appointment Extract – Filtering".
        """
        category_lower = category.lower()

        for pattern, mapping in CATEGORY_TO_CONFIG_KEYS.items():
            if pattern in category_lower:
                return mapping

        return None

    def _slugify(self, text: str) -> str:
        """Convert text to a slug for unmapped categories."""
        # Remove special chars, convert spaces to underscores
        slug = re.sub(r'[^\w\s]', '', text.lower())
        slug = re.sub(r'\s+', '_', slug.strip())
        return slug[:50]  # Limit length

    # =========================================================================
    # SPECIALIZED PARSERS - Parse complex cell values
    # =========================================================================

    def _parse_appointment_extract(self, override: str, default: str) -> Dict[str, str]:
        """
        Parse Patient Appointment Extract – Filtering cell.

        PURPOSE: Extract separate config values from a cell containing:
                 - Patient Status: New
                 - Providers: ... (by location)

        PROVIDER FORMAT: Location-specific providers are listed as:
            Location Specific Filters:
            PCI Breast Surgery West:
            Providers:
            - Jessica Bautista,
            - Rachel Dise
            PCI Franz Providers:
            - Christine Kemp
            - Nora Lersch

        This should produce:
            extract_providers@PCI BREAST SURGERY WEST: "Jessica Bautista, Rachel Dise"
            extract_providers@PCI FRANZ BREAST CARE: "Christine Kemp, Nora Lersch"

        RETURNS:
            Dict mapping config_key to value, with location-specific keys using @ suffix
        """
        result = {}
        text = override or default
        if not text:
            return result

        lines = text.split('\n')
        locations = self.result.get('scope_locations', [])

        # First pass: extract patient_status
        for line in lines:
            line = line.strip()
            if not line:
                continue

            line_lower = line.lower()
            if line_lower.startswith('patient status'):
                if ':' in line:
                    value = line.split(':', 1)[1].strip()
                    if value:
                        result['extract_patient_status'] = value
                break

        # Second pass: extract location-specific providers
        # Look for "Location Specific Filters:" section
        in_provider_section = False
        current_location = None
        current_providers = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            line_lower = line.lower()

            # Check for start of location-specific section
            if 'location specific' in line_lower:
                in_provider_section = True
                continue

            if not in_provider_section:
                continue

            # Check if this line starts a new location
            # Patterns: "PCI Breast Surgery West:" or "PCI Franz Providers:"
            matched_location = None

            # Try to match this line as a location header
            if ':' in line:
                before_colon = line.split(':')[0].strip()
                # Remove "Providers" suffix if present
                before_colon = re.sub(r'\s*providers?\s*$', '', before_colon, flags=re.IGNORECASE)

                matched_location = self._match_location_name(before_colon, locations)

            # Also check for location name without colon on its own line
            if not matched_location and not line.startswith('-'):
                test_name = re.sub(r'\s*providers?\s*$', '', line, flags=re.IGNORECASE)
                matched_location = self._match_location_name(test_name, locations)

            if matched_location:
                # Save previous location's providers
                if current_location and current_providers:
                    provider_str = ', '.join(current_providers)
                    result[f'extract_providers@{current_location}'] = provider_str

                current_location = matched_location
                current_providers = []
                continue

            # Check if this is a provider line (bullet point)
            if line.startswith('-') or line.startswith('•'):
                provider_name = line.lstrip('-•').strip()
                # Clean up trailing commas
                provider_name = provider_name.rstrip(',').strip()
                if provider_name and current_location:
                    current_providers.append(provider_name)

        # Save last location's providers
        if current_location and current_providers:
            provider_str = ', '.join(current_providers)
            result[f'extract_providers@{current_location}'] = provider_str

        return result

    def _parse_helpdesk_email(self, override: str, default: str) -> Dict[str, str]:
        """
        Parse Help Desk Email cell.

        PURPOSE: Separate actual email address from workflow description.

        LOGIC:
        - If value contains "@", extract that as helpdesk_email
        - If value describes a process/workflow, store as helpdesk_workflow
        - Both can be present in one cell

        RETURNS:
            Dict with helpdesk_email and/or helpdesk_workflow
        """
        result = {}
        text = override or default
        if not text:
            return result

        # Look for email addresses
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        emails = re.findall(email_pattern, text)

        if emails:
            # Found email(s) - use the first one
            result['helpdesk_email'] = emails[0]

        # Check if there's workflow description
        # Workflow indicators: "respond", "workflow", "process", "guidance"
        workflow_indicators = ['respond', 'workflow', 'process', 'guidance',
                               'contact', 'follow up', 'escalate']
        text_lower = text.lower()

        if any(indicator in text_lower for indicator in workflow_indicators):
            # Store the full text as workflow (it describes a process)
            result['helpdesk_workflow'] = text

        # If no email found but text exists, it might all be workflow
        if not emails and text:
            result['helpdesk_workflow'] = text

        return result

    def _parse_location_specific_value(self, override: str, default: str,
                                        config_keys: List[str]) -> Dict[str, str]:
        """
        Parse location-specific values from a cell.

        PURPOSE: Handle cells with format like:
                 "Breast Surgery West: 503.216.6407
                  Franz Breast Care: 503.215.7920"

                 OR shared values:
                 "Location A, Location B: shared_value
                  Location C: different_value"

        For location-specific values, we use fuzzy matching to map
        document location names to actual scope_locations from the Scope section.

        RETURNS:
            Dict mapping config_key to value, with location-specific keys
            in format: config_key@location_name
        """
        result = {}
        text = override or default
        if not text:
            return result

        if not config_keys:
            return result

        config_key = config_keys[0]

        # Use the new distribution method
        distributed = self._distribute_to_locations(text, self.result.get('scope_locations', []))

        if distributed:
            # Store location-specific values with @location suffix
            # This format tells the importer to route to location-level configs
            for location_name, value in distributed.items():
                result[f'{config_key}@{location_name}'] = value

            # Also store the raw location map for reference
            result[f'{config_key}_by_location'] = distributed
        else:
            # Not location-specific - simple value
            result[config_key] = text

        return result

    def _parse_location_specific_hours(self, override: str, default: str,
                                        config_keys: List[str]) -> Dict[str, str]:
        """
        Parse location-specific hours of operation.

        PURPOSE: Handle cells with format like:
                 "Breast Surgery West: 8am-5pm
                  Franz Breast Care: 9am-6pm"

        Parses both the location distribution AND the time ranges.

        RETURNS:
            Dict with hours_open@location and hours_close@location keys
        """
        result = {}
        text = override or default
        if not text:
            return result

        # First distribute by location
        distributed = self._distribute_to_locations(text, self.result.get('scope_locations', []))

        if distributed:
            # Parse hours for each location
            for location_name, hours_text in distributed.items():
                parsed_hours = self._parse_hours(hours_text, '')

                if 'hours_open' in parsed_hours:
                    result[f'hours_open@{location_name}'] = parsed_hours['hours_open']
                if 'hours_close' in parsed_hours:
                    result[f'hours_close@{location_name}'] = parsed_hours['hours_close']

            # Store raw for reference
            result['hours_by_location'] = distributed
        else:
            # Not location-specific - parse as simple hours
            parsed = self._parse_hours(text, '')
            result.update(parsed)

        return result

    def build_location_fragments(self, location_name: str) -> List[str]:
        """
        Build searchable fragments from a location name for fuzzy matching.

        PURPOSE: Generate multiple text fragments from a canonical location name
                 so we can detect when that location is referenced in config values.

        PARAMETERS:
            location_name: Full location name from Scope section
                           e.g., "PCI BREAST SURGERY WEST (Includes PCI BREAST CARE CLINIC WEST)"

        RETURNS:
            List of lowercase fragments for matching
            e.g., ["breast surgery west", "breast surgery", "breast care clinic west", "breast care"]

        EXAMPLES:
            Input: "PCI BREAST SURGERY WEST (Includes PCI BREAST CARE CLINIC WEST)"
            Output: ["breast surgery west", "breast surgery", "breast care clinic west",
                     "breast care", "bsw"]

            Input: "PCI FRANZ BREAST CARE"
            Output: ["franz breast care", "franz breast", "franz"]

            Input: "MONTANA CLINIC A"
            Output: ["montana clinic a", "montana clinic", "montana"]

        WHY THIS APPROACH: Document authors use abbreviated names like "Franz" or
        "Breast Surgery West" instead of full names. We need multiple fragments
        to catch these variations.
        """
        fragments = []

        # Clean the name - remove common prefixes and parenthetical content
        clean_name = location_name.lower()

        # Remove common prefixes
        prefixes_to_remove = ['pci ', 'prov ', 'oph ', 'owf ', 'o ']
        for prefix in prefixes_to_remove:
            if clean_name.startswith(prefix):
                clean_name = clean_name[len(prefix):]

        # Extract parenthetical content separately (e.g., "Includes PCI BREAST CARE CLINIC WEST")
        paren_match = re.search(r'\(([^)]+)\)', clean_name)
        paren_content = None
        if paren_match:
            paren_content = paren_match.group(1).lower()
            # Remove parenthetical from main name
            clean_name = re.sub(r'\s*\([^)]+\)\s*', '', clean_name).strip()
            # Clean the parenthetical content too
            for prefix in prefixes_to_remove:
                if paren_content.startswith(prefix):
                    paren_content = paren_content[len(prefix):]
            paren_content = re.sub(r'^includes?\s+', '', paren_content).strip()

        # Add the full clean name
        if clean_name:
            fragments.append(clean_name.strip())

        # Add parenthetical content fragments
        if paren_content:
            fragments.append(paren_content.strip())

        # Generate progressively shorter fragments
        # "breast surgery west" → "breast surgery", "surgery west"
        for name in [clean_name, paren_content]:
            if not name:
                continue
            words = name.split()
            # Add 2-word combinations
            if len(words) >= 2:
                for i in range(len(words) - 1):
                    fragment = ' '.join(words[i:i+2])
                    if fragment not in fragments and len(fragment) > 3:
                        fragments.append(fragment)
            # Add single significant words (not common words)
            skip_words = {'the', 'and', 'of', 'at', 'in', 'clinic', 'center',
                          'care', 'surgery', 'breast', 'includes', 'pci'}
            for word in words:
                if word not in skip_words and len(word) > 2 and word not in fragments:
                    fragments.append(word)

        # Add abbreviation (first letter of each significant word)
        words = clean_name.split()
        if len(words) >= 2:
            abbrev = ''.join(w[0] for w in words if w not in {'the', 'and', 'of', 'at'})
            if len(abbrev) >= 2:
                fragments.append(abbrev)

        return fragments

    def is_location_specific(self, value_text: str, locations: List[str]) -> bool:
        """
        Determine if a value contains location-specific data.

        PURPOSE: Check if the value text references ANY of the known location names.
                 If yes, it needs to be parsed and distributed.
                 If no, it's a shared value that applies to all locations.

        PARAMETERS:
            value_text: The config value text from the document
            locations: List of location names from Scope section (dynamic, not hardcoded)

        RETURNS:
            True if value contains location references, False if it's a shared value

        EXAMPLES:
            "Ambry" → False (no location references, shared value)
            "Breast Surgery West: 503.555.1234" → True (contains location reference)
            "Same as default" → False (shared value)

        WHY THIS APPROACH: We must detect location-specific values generically
        without hardcoding location names. Works for Portland, Montana, or any future clinic.
        """
        if not value_text or not locations:
            return False

        value_lower = value_text.lower()

        # Build fragments for each location and check if any appear in the value
        for location in locations:
            fragments = self.build_location_fragments(location)
            for fragment in fragments:
                if fragment in value_lower:
                    return True

        return False

    def _distribute_to_locations(self, value_text: str, locations: List[str]) -> Dict[str, str]:
        """
        Parse a value and distribute to locations - THE CORE DISTRIBUTION METHOD.

        PURPOSE: This is the main entry point for distributing config values.
                 Every config value must be distributed to locations:
                 - If location-specific: parse and route to matching locations
                 - If shared: apply the same value to ALL locations

        PARAMETERS:
            value_text: Text that may contain location-specific values
            locations: List of known location names from Scope section (dynamic)

        RETURNS:
            Dict mapping location names to their values.
            NEVER returns empty dict - always returns values for locations.

        ALGORITHM:
            1. Skip empty values or "same as default"
            2. Check if value is location-specific using is_location_specific()
            3. If location-specific: parse and distribute to specific locations
            4. If shared: distribute same value to ALL locations

        R EQUIVALENT: Like case_when() with grepl() for pattern matching,
                      combined with fuzzy matching via agrep()

        WHY THIS APPROACH: We must handle ANY program's locations generically.
        Portland has "PCI Breast Surgery West", Montana might have "Billings Center".
        The parser must work for all without hardcoding.
        """
        if not value_text or not locations:
            return {}

        # Skip "same as default" values
        if 'same as default' in value_text.lower():
            return {}

        # Check if this value contains location-specific data
        if self.is_location_specific(value_text, locations):
            # Parse location-specific values
            return self._parse_location_specific_values(value_text, locations)
        else:
            # Shared value - distribute to ALL locations
            return {loc: value_text for loc in locations}

    def _parse_location_specific_values(self, value_text: str, locations: List[str]) -> Dict[str, str]:
        """
        Parse text that contains location-specific values.

        PURPOSE: Extract location-value pairs from text that references specific locations.

        HANDLES FORMATS:
            FORMAT 1: "Location A: value1\\nLocation B: value2"
            FORMAT 2: "Location A, Location B, Location C: shared_value\\nLocation D: different_value"
            FORMAT 3: "All: shared_value" (distributes to all locations)
            FORMAT 4: "Location A Providers: name1, name2\\nLocation B Provider: name3"
            FORMAT 5: "Location A:\\n- value1\\nLocation B:\\n- value2" (bullet format)
            FORMAT 6: "Location A:\\nvalue1\\nLocation B:\\nvalue2" (value on next line, no bullet)

        PARAMETERS:
            value_text: Text containing location-specific values
            locations: List of known location names from Scope section

        RETURNS:
            Dict mapping matched location names to their values.

        WHY THIS APPROACH: Document authors use various formats for location-specific data.
        We must handle all common patterns while using fuzzy matching for location names.
        Common formats include value on next line (with or without bullet).
        """
        if not value_text or not locations:
            return {}

        result = {}
        lines = value_text.strip().split('\n')

        current_locations = []  # Locations waiting for a value on the next line
        i = 0

        while i < len(lines):
            line = lines[i].strip()
            i += 1

            if not line:
                continue

            # Skip URLs
            if line.lower().startswith('http'):
                continue

            # Check if this line is a bullet point value (starts with -)
            # This handles FORMAT 5: "Location A:\n- value"
            if line.startswith('-') or line.startswith('•'):
                value = line.lstrip('-•').strip()
                if value and current_locations:
                    # Apply this value to all pending locations
                    for loc in current_locations:
                        result[loc] = value
                    current_locations = []
                continue

            # Check for "Location: Value" format
            if ':' not in line:
                # This is a value line for pending locations (FORMAT 6)
                # Lines without colons are values, not location headers
                if current_locations:
                    # This is the value for the pending locations
                    for loc in current_locations:
                        result[loc] = line
                    current_locations = []
                continue

            parts = line.split(':', 1)
            if len(parts) != 2:
                continue

            location_part = parts[0].strip()
            value = parts[1].strip()

            if not location_part:
                continue

            # Check for "All" or "All Locations" meaning distribute to all
            if location_part.lower() in ['all', 'all locations', 'default']:
                if value:
                    for loc in locations:
                        result[loc] = value
                else:
                    # Value might be on next line
                    current_locations = list(locations)
                continue

            # Parse location names (comma-separated or single)
            if ',' in location_part:
                location_names = [name.strip() for name in location_part.split(',')]
            else:
                location_names = [location_part]

            # Match each location name to scope_locations using fuzzy matching
            matched_locations = []
            for loc_name in location_names:
                matched = self._match_location_name(loc_name, locations)
                if matched:
                    matched_locations.append(matched)

            if matched_locations:
                if value:
                    # Value is on same line - apply immediately
                    for loc in matched_locations:
                        result[loc] = value
                else:
                    # Value might be on next line (FORMAT 5 or 6)
                    current_locations = matched_locations

        return result

    def _looks_like_location_reference(self, text: str, locations: List[str]) -> bool:
        """
        Check if text looks like it references a location name.

        PURPOSE: Distinguish between value text and location reference text
                 when parsing multi-line location-specific values.

        PARAMETERS:
            text: The text to check
            locations: List of known location names

        RETURNS:
            True if text appears to reference a location, False otherwise

        WHY THIS APPROACH: When parsing "Location:\nValue\nLocation2:\nValue2",
        we need to know if a line without a colon is a value or a partial location name.
        """
        if not text or not locations:
            return False

        text_lower = text.lower()

        # Check if any location fragment appears in the text
        for location in locations:
            fragments = self.build_location_fragments(location)
            for fragment in fragments[:3]:  # Check first few fragments
                if fragment in text_lower:
                    return True

        return False

    def _match_location_name(self, text_name: str, location_names: List[str]) -> Optional[str]:
        """
        Fuzzy match location names from text to actual location names.

        PURPOSE: Match document location names to scope_locations,
                 handling variations like:
                 - "Breast Surgery West" → "PCI BREAST SURGERY WEST"
                 - "Franz" → "PCI Franz Breast Care"
                 - "BSW" → "PCI Breast Surgery West" (abbreviation)

        PARAMETERS:
            text_name: Location name as it appears in document
            location_names: List of canonical location names from Scope

        RETURNS:
            The matched location name from location_names, or None if no match

        R EQUIVALENT: Like agrep() or stringdist::stringdistmatrix() for fuzzy matching

        WHY THIS APPROACH: Document authors may use abbreviations or informal
        names that don't exactly match the canonical location names.
        """
        if not text_name or not location_names:
            return None

        text_lower = text_name.lower().strip()

        # First try exact match (case-insensitive)
        for loc in location_names:
            if loc.lower() == text_lower:
                return loc

        # Second: check if text is contained in any location name
        # e.g., "Breast Surgery West" in "PCI BREAST SURGERY WEST"
        for loc in location_names:
            loc_lower = loc.lower()
            if text_lower in loc_lower:
                return loc
            # Also check reverse: "Franz Breast Care" might be in text
            if loc_lower in text_lower:
                return loc

        # Third: check for word overlap (partial match)
        # Split both names into words and check overlap
        text_words = set(re.findall(r'\w+', text_lower))

        best_match = None
        best_score = 0

        for loc in location_names:
            loc_words = set(re.findall(r'\w+', loc.lower()))

            # Remove common words that don't help matching
            skip_words = {'pci', 'the', 'and', 'of', 'at', 'center', 'clinic'}
            text_significant = text_words - skip_words
            loc_significant = loc_words - skip_words

            if not text_significant or not loc_significant:
                continue

            # Calculate overlap score
            overlap = text_significant & loc_significant
            score = len(overlap) / max(len(text_significant), len(loc_significant))

            if score > best_score and score >= 0.5:  # At least 50% word overlap
                best_score = score
                best_match = loc

        if best_match:
            return best_match

        # Fourth: check for abbreviation patterns
        # e.g., "BSW" for "Breast Surgery West"
        if len(text_name) <= 4 and text_name.isupper():
            # Might be an abbreviation - check initials
            for loc in location_names:
                # Get initials from location name
                words = re.findall(r'[A-Z][a-z]*|\b\w+', loc)
                initials = ''.join(w[0].upper() for w in words if w and w[0].isalpha())

                # Also try without 'PCI' prefix
                loc_no_prefix = re.sub(r'^PCI\s+', '', loc, flags=re.IGNORECASE)
                words_no_prefix = re.findall(r'[A-Z][a-z]*|\b\w+', loc_no_prefix)
                initials_no_prefix = ''.join(w[0].upper() for w in words_no_prefix if w and w[0].isalpha())

                if text_name.upper() in [initials, initials_no_prefix]:
                    return loc

        return None

    def _parse_hours(self, override: str, default: str) -> Dict[str, str]:
        """
        Parse hours of operation.

        PURPOSE: Extract open/close times from text like "8am-5pm" or "8:00 AM to 5:00 PM"

        RETURNS:
            Dict with hours_open and hours_close
        """
        result = {}
        text = override or default
        if not text:
            return result

        # Pattern for time ranges
        time_pattern = r'(\d{1,2}(?::\d{2})?\s*(?:am|pm)?)\s*[-–to]+\s*(\d{1,2}(?::\d{2})?\s*(?:am|pm)?)'
        match = re.search(time_pattern, text, re.IGNORECASE)

        if match:
            result['hours_open'] = match.group(1).strip()
            result['hours_close'] = match.group(2).strip()
        else:
            # Store raw if can't parse
            result['hours_open'] = text

        return result

    def _parse_age_range(self, override: str, default: str) -> Dict[str, str]:
        """
        Parse TC age range.

        PURPOSE: Extract minimum and maximum ages from text like "18-75" or "18 to 75 years"

        RETURNS:
            Dict with tc_minimum_age and tc_maximum_age
        """
        result = {}
        text = override or default
        if not text:
            return result

        # Pattern for age ranges
        age_pattern = r'(\d+)\s*[-–to]+\s*(\d+)'
        match = re.search(age_pattern, text)

        if match:
            result['tc_minimum_age'] = match.group(1)
            result['tc_maximum_age'] = match.group(2)
        else:
            # Try to find just minimum age
            min_pattern = r'(?:min(?:imum)?|>=?)\s*(\d+)'
            min_match = re.search(min_pattern, text, re.IGNORECASE)
            if min_match:
                result['tc_minimum_age'] = min_match.group(1)

        return result

    def _parse_lab_order(self, override: str, default: str) -> Dict[str, str]:
        """
        Parse lab order information.

        PURPOSE: Extract test code and test name from lab order text.

        RETURNS:
            Dict with lab_default_test_code and lab_default_test_name
        """
        result = {}
        text = override or default
        if not text:
            return result

        # Look for test code pattern (numeric)
        code_pattern = r'(?:test\s*code|code)[:\s]*(\d+)'
        code_match = re.search(code_pattern, text, re.IGNORECASE)
        if code_match:
            result['lab_default_test_code'] = code_match.group(1)

        # Look for test name pattern
        name_pattern = r'(?:test\s*name|name)[:\s]*([^\n]+)'
        name_match = re.search(name_pattern, text, re.IGNORECASE)
        if name_match:
            result['lab_default_test_name'] = name_match.group(1).strip()

        # If no structured format, try to split by newline
        if not result:
            lines = [l.strip() for l in text.split('\n') if l.strip()]
            if lines:
                # First numeric-looking thing is probably code
                for line in lines:
                    if re.match(r'^\d+$', line):
                        result['lab_default_test_code'] = line
                    elif not line.startswith('Test'):
                        result['lab_default_test_name'] = line
                        break

        return result

    # =========================================================================
    # HEADER AND METADATA PARSING
    # =========================================================================

    def _parse_header(self) -> None:
        """
        Extract Doc ID, Version, Parent SRS from header paragraphs.

        WHY THIS APPROACH: The first few paragraphs typically contain
        document metadata in a semi-structured format. We look for
        patterns like "Doc ID: XXX" or "Version: X.X".
        """
        patterns = {
            'doc_id': [
                r'Doc(?:ument)?\s*ID[:\s]+([A-Z0-9\-]+)',
                r'Document\s+ID[:\s]+([A-Z0-9\-]+)',
                r'^([A-Z0-9]+-[A-Z]+-[A-Z]+-[A-Z]+)$'
            ],
            'version': [
                r'Version[:\s]+([0-9]+\.[0-9]+)',
                r'v([0-9]+\.[0-9]+)',
                r'Ver[:\s]+([0-9]+\.[0-9]+)'
            ],
            'parent_srs': [
                r'Parent\s+SRS[:\s]+([A-Z0-9\-_\.]+)',
                r'SRS\s+Reference[:\s]+([A-Z0-9\-_\.]+)',
                r'Based\s+on[:\s]+([A-Z0-9\-_\.]+)'
            ]
        }

        for i, para in enumerate(self.doc.paragraphs[:20]):
            text = para.text.strip()
            if not text:
                continue

            for field, field_patterns in patterns.items():
                if self.result[field]:
                    continue

                for pattern in field_patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        self.result[field] = match.group(1)
                        break

    def _extract_clinic_name(self) -> None:
        """
        Extract clinic name from document ID or content.

        WHY THIS APPROACH: Clinic name is often embedded in the doc_id
        (e.g., P4M-CL-PORT-SPEC → Portland) or in the title.
        """
        abbrev_map = {
            'PORT': 'Portland',
            'PROV': 'Providence',
            'SEA': 'Seattle',
            'PDX': 'Portland',
            'LA': 'Los Angeles',
            'SF': 'San Francisco',
            'CHI': 'Chicago',
            'NYC': 'New York'
        }

        if self.result['doc_id']:
            match = re.search(r'-CL-([A-Z]+)-', self.result['doc_id'])
            if match:
                abbrev = match.group(1)
                self.result['clinic_name'] = abbrev_map.get(abbrev, abbrev.title())
                return

        for para in self.doc.paragraphs[:10]:
            text = para.text.strip()
            if para.style and 'Heading' in para.style.name:
                match = re.search(r'(\w+)\s+Clinic', text, re.IGNORECASE)
                if match:
                    self.result['clinic_name'] = match.group(1)
                    return

        for para in self.doc.paragraphs[:15]:
            text = para.text.strip()
            for program in ['Prevention4ME', 'Precision4ME', 'GenoRx', 'Discover']:
                if program.lower() in text.lower():
                    self.result['program'] = program
                    break

    def _parse_scope(self) -> None:
        """
        Extract location names from Scope section.

        WHY THIS APPROACH: The Scope section typically lists which
        locations/service points this specification covers.
        """
        in_scope = False
        scope_lines = []

        for para in self.doc.paragraphs:
            text = para.text.strip()

            if 'scope' in text.lower() and len(text) < 50:
                in_scope = True
                continue

            if in_scope:
                if para.style and 'Heading' in para.style.name and 'scope' not in text.lower():
                    break

                if text and not text.startswith('•'):
                    if self._looks_like_location(text):
                        scope_lines.append(text)

                if text.startswith('•') or text.startswith('-'):
                    location = text.lstrip('•-').strip()
                    if self._looks_like_location(location):
                        scope_lines.append(location)

        self.result['scope_locations'] = scope_lines

    def _looks_like_location(self, text: str) -> bool:
        """Check if text looks like a location name."""
        skip_patterns = [
            r'^this\s+document',
            r'^the\s+following',
            r'^applies\s+to',
            r'^version',
            r'^date',
            r'^\d+\.\d+'
        ]

        for pattern in skip_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return False

        location_indicators = [
            'surgery', 'clinic', 'center', 'care', 'hospital',
            'medical', 'health', 'west', 'east', 'north', 'south'
        ]

        text_lower = text.lower()
        return any(indicator in text_lower for indicator in location_indicators)

    # =========================================================================
    # TABLE PARSING
    # =========================================================================

    def _parse_all_tables(self) -> None:
        """
        Parse all tables in the document.

        WHY THIS APPROACH: We identify table type by its headers
        and route to appropriate parsing method.
        """
        for table in self.doc.tables:
            if len(table.rows) < 2:
                continue

            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

            self.result['raw_tables'].append({
                'headers': headers,
                'row_count': len(table.rows)
            })

            if self._is_config_table(headers):
                self._parse_config_table(table)
            elif self._is_change_log_table(headers):
                self._parse_change_log(table)
            elif self._is_provider_table(headers):
                self._parse_provider_table(table)

    def _is_config_table(self, headers: List[str]) -> bool:
        """Check if table headers indicate a configuration table."""
        config_indicators = ['category', 'global', 'default', 'override', 'customization']
        matches = sum(1 for h in headers for ind in config_indicators if ind in h)
        return matches >= 2

    def _is_change_log_table(self, headers: List[str]) -> bool:
        """Check if table headers indicate a change log table."""
        log_indicators = ['date', 'version', 'author', 'summary', 'change', 'description']
        matches = sum(1 for h in headers for ind in log_indicators if ind in h)
        return matches >= 2

    def _is_provider_table(self, headers: List[str]) -> bool:
        """Check if table headers indicate a provider table."""
        provider_indicators = ['provider', 'npi', 'name', 'role', 'ordering']
        matches = sum(1 for h in headers for ind in provider_indicators if ind in h)
        return matches >= 2

    def _parse_config_table(self, table: Table) -> None:
        """
        Parse the main configuration matrix table.

        PURPOSE: Extract configuration settings from the config table
        """
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        col_map = self._map_config_columns(headers)

        if not col_map:
            return

        for row in table.rows[1:]:
            cells = row.cells

            if len(cells) < 2:
                continue

            category = self._clean_cell(cells[col_map.get('category', 0)].text) if col_map.get('category') is not None else ''
            global_default = self._clean_cell(cells[col_map.get('global', 1)].text) if col_map.get('global') is not None else ''
            raw_override = cells[col_map.get('override', 2)].text.strip() if col_map.get('override') is not None else ''
            rationale = self._clean_cell(cells[col_map.get('rationale', 3)].text) if col_map.get('rationale') is not None else ''

            if not category and not global_default and not raw_override:
                continue

            # Check if this is the provider row - parse providers from cell
            if 'provider' in category.lower():
                providers = self._parse_providers_from_cell(raw_override)
                self.result['providers'].extend(providers)
                continue

            # Parse complex values
            parsed_override = self._parse_complex_value(raw_override) if raw_override else None

            config = {
                'category': category,
                'global_default': global_default,
                'override': raw_override,
                'parsed_override': parsed_override,
                'rationale': rationale,
                'raw_headers': headers
            }

            self.result['configurations'].append(config)

    def _map_config_columns(self, headers: List[str]) -> Dict[str, int]:
        """Map header names to column indices."""
        col_map = {}

        for i, header in enumerate(headers):
            h_lower = header.lower()

            if 'category' in h_lower or 'setting' in h_lower:
                col_map['category'] = i
            elif 'global' in h_lower or 'default' in h_lower:
                col_map['global'] = i
            elif 'override' in h_lower or 'custom' in h_lower or self.result.get('clinic_name', '').lower() in h_lower:
                col_map['override'] = i
            elif 'rationale' in h_lower or 'source' in h_lower or 'notes' in h_lower:
                col_map['rationale'] = i

        if 'override' not in col_map:
            for i, header in enumerate(headers):
                if i not in col_map.values() and len(header) > 2:
                    col_map['override'] = i
                    break

        return col_map

    def _parse_change_log(self, table: Table) -> None:
        """Parse the change log table."""
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

        col_map = {}
        for i, h in enumerate(headers):
            if 'date' in h:
                col_map['date'] = i
            elif 'version' in h or 'ver' in h:
                col_map['version'] = i
            elif 'author' in h or 'by' in h:
                col_map['author'] = i
            elif 'summary' in h or 'description' in h or 'change' in h:
                col_map['summary'] = i

        for row in table.rows[1:]:
            cells = row.cells

            entry = {
                'date': self._clean_cell(cells[col_map.get('date', 0)].text) if col_map.get('date') is not None else '',
                'version': self._clean_cell(cells[col_map.get('version', 1)].text) if col_map.get('version') is not None else '',
                'author': self._clean_cell(cells[col_map.get('author', 2)].text) if col_map.get('author') is not None else '',
                'summary': self._clean_cell(cells[col_map.get('summary', 3)].text) if col_map.get('summary') is not None else ''
            }

            if entry['date'] or entry['version'] or entry['summary']:
                self.result['change_log'].append(entry)

    def _parse_provider_table(self, table: Table) -> None:
        """Parse provider information table."""
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

        col_map = {}
        for i, h in enumerate(headers):
            if 'provider' in h or 'name' in h:
                col_map['name'] = i
            elif 'npi' in h:
                col_map['npi'] = i
            elif 'location' in h or 'site' in h:
                col_map['location'] = i
            elif 'role' in h:
                col_map['role'] = i

        for row in table.rows[1:]:
            cells = row.cells

            provider = {
                'name': self._clean_cell(cells[col_map.get('name', 0)].text) if col_map.get('name') is not None else '',
                'npi': self._clean_cell(cells[col_map.get('npi', 1)].text) if col_map.get('npi') is not None else '',
                'location': self._clean_cell(cells[col_map.get('location', 2)].text) if col_map.get('location') is not None else '',
                'role': self._clean_cell(cells[col_map.get('role', 3)].text) if col_map.get('role') is not None else 'Ordering Provider'
            }

            if provider['name']:
                self.result['providers'].append(provider)

    def _parse_complex_value(self, value_text: str) -> Optional[Dict]:
        """
        Parse complex values like location-specific values.

        EXAMPLE INPUT:
            'Breast Surgery West: 503.216.6407
             Franz Breast Care: 503.215.7920'

        RETURNS:
            {
                'Breast Surgery West': '503.216.6407',
                'Franz Breast Care': '503.215.7920'
            }
        """
        if not value_text:
            return None

        if ':' not in value_text:
            return None

        result = {}
        lines = value_text.split('\n')

        for line in lines:
            line = line.strip()
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    key = parts[0].strip()
                    val = parts[1].strip()
                    if key and val:
                        result[key] = val

        return result if result else None

    def _parse_providers_from_cell(self, cell_text: str) -> List[Dict]:
        """
        Parse providers from a cell containing a list of providers by location.

        PURPOSE: Extract provider names and NPIs from a multi-line cell

        EXPECTED FORMAT:
            PROV PCI Breast Surgery West
            Jessica Bautista,, NP
            (NPI 1184485492)
            OPH PCI Franz Breast Care
            Christine Kemp, NP
            (NPI 1215158639)

        RETURNS:
            List of provider dicts with: name, npi, location, role
        """
        providers = []
        current_location = None

        if not cell_text:
            return providers

        lines = cell_text.strip().split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            if self._is_location_header(line):
                current_location = self._extract_location_name(line)
                i += 1
                continue

            credential_match = re.search(r',?\s*(MD|NP|PA|DO|RN|ARNP)\s*$', line, re.IGNORECASE)
            if credential_match:
                provider_name = line.strip()
                provider_name = re.sub(r',+', ',', provider_name)
                provider_name = re.sub(r'\s+', ' ', provider_name)

                npi = None
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    npi_match = re.search(r'\(?\s*NPI\s*[:\s]*(\d{9,10})\s*\)?', next_line, re.IGNORECASE)
                    if npi_match:
                        npi = npi_match.group(1)
                        i += 1

                providers.append({
                    'name': provider_name,
                    'npi': npi,
                    'location': current_location or 'Unknown',
                    'role': 'Ordering Provider'
                })

            i += 1

        return providers

    def _is_location_header(self, line: str) -> bool:
        """Check if a line is a location header."""
        prefixes = ['PROV ', 'OPH ', 'OWF ', 'O ']
        for prefix in prefixes:
            if line.upper().startswith(prefix):
                return True

        if 'PCI' in line.upper() and not re.search(r'(MD|NP|PA|DO)\s*$', line, re.IGNORECASE):
            return True

        return False

    def _extract_location_name(self, line: str) -> str:
        """Extract clean location name from a location header line."""
        result = line
        for prefix in ['PROV ', 'OPH ', 'OWF ', 'O ']:
            if result.upper().startswith(prefix):
                result = result[len(prefix):]
                break

        return result.strip()

    def _clean_cell(self, text: str) -> str:
        """Clean cell text by removing extra whitespace."""
        if not text:
            return ''

        text = re.sub(r'[\s\u00a0]+', ' ', text)
        text = text.strip()

        return text


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def parse_clinic_spec(file_path: str) -> Dict:
    """
    Quick function to parse a clinic spec document.

    PARAMETERS:
        file_path: Path to .docx file

    RETURNS:
        Parsed document data as dict
    """
    parser = ClinicSpecParser(file_path)
    return parser.parse()


# =============================================================================
# MODULE TEST
# =============================================================================

if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 2:
        print("Usage: python word_parser.py <path_to_doc.docx>")
        print("\nThis module parses clinic specification Word documents.")
        print("Provide a .docx file path to test the parser.")
        sys.exit(1)

    file_path = sys.argv[1]
    print(f"Parsing: {file_path}")

    parser = ClinicSpecParser(file_path)
    result = parser.parse()

    print("\n" + "=" * 60)
    print("PARSE RESULTS")
    print("=" * 60)

    print(f"\nDoc ID: {result['doc_id']}")
    print(f"Version: {result['version']}")
    print(f"Clinic Name: {result['clinic_name']}")

    print(f"\nScope Locations ({len(result['scope_locations'])}):")
    for loc in result['scope_locations']:
        print(f"  - {loc}")

    print(f"\nRaw Configurations ({len(result['configurations'])}):")
    for config in result['configurations'][:5]:
        print(f"  - {config['category'][:50]}...")

    print(f"\nMapped Configurations ({len(result['mapped_configs'])}):")
    for config in result['mapped_configs']:
        print(f"  - {config['config_key']}: {str(config['value'])[:50]}...")

    print(f"\nProviders ({len(result['providers'])}):")
    for prov in result['providers']:
        print(f"  - {prov['name']} (NPI: {prov.get('npi', 'N/A')}) @ {prov.get('location', 'Unknown')}")

    print(f"\nTables found: {len(result['raw_tables'])}")
