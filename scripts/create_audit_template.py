#!/usr/bin/env python3
"""
Generate the audit completion memo template with Jinja2 placeholders.
Run once to create the template file.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_template():
    doc = Document()

    # Title
    title = doc.add_heading('INTERNAL COMPLIANCE RECORD', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Access Review Completion Memo', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header info
    doc.add_paragraph()
    doc.add_paragraph("System: Propel Health Platform")
    doc.add_paragraph("Client: {{clinic_name}}")
    doc.add_paragraph("Audit Cycle: {{audit_year}} Access Review -- {{audit_type}}")
    doc.add_paragraph("Document Version: {{document_version}}")
    doc.add_paragraph("Date Completed: {{date_confirmed}}")

    # Section 1: Overview
    doc.add_heading('1. Overview', level=2)
    doc.add_paragraph(
        'This memo formally documents the completion of the {{audit_year}} access review '
        'for the {{clinic_name}} as required under:'
    )

    # Bullet list for regulations
    doc.add_paragraph('21 CFR Part 11 – Section 11.10(i) (Limiting system access to authorized individuals)', style='List Bullet')
    doc.add_paragraph('HIPAA 45 CFR §164.308(a)(4) (Information access management)', style='List Bullet')
    doc.add_paragraph('SOC 2 – CC6.1, CC6.2, CC6.3 (Access controls and user authorizations)', style='List Bullet')

    doc.add_paragraph(
        'The review was conducted in accordance with the Propel Health Platform '
        'Critical Infrastructure Access Control Policy, version 1.0, approved February 28, 2025.'
    )

    # Section 2: Audit Timeline
    doc.add_heading('2. Audit Timeline', level=2)

    # Timeline table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Date'

    # Data rows with placeholders
    timeline_data = [
        ('Initiated by PHP', 'Current access roster submitted for review to Clinic Manager', '{{date_initiated}}'),
        ('Reviewed by Client Program Liaison', 'Access validation returned to Propel by Clinic Manager', '{{date_reviewed}}'),
        ('Internal Review Finalized', 'Findings validated by Compliance & Ops team', '{{date_finalized}}'),
        ('Change Tickets Submitted', 'Revocation/Modification/Provision requests (#{{ticket_number}}) sent to Dev Team', '{{date_tickets_submitted}}'),
        ('Final Confirmation', 'Update confirmed by Development Team via Zendesk ticket.', '{{date_confirmed}}'),
    ]

    for i, (step, desc, date) in enumerate(timeline_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = step
        row_cells[1].text = desc
        row_cells[2].text = date

    # Section 3: Summary of Findings
    doc.add_heading('3. Summary of Findings', level=2)
    doc.add_paragraph("Total Accounts Reviewed: {{total_reviewed}}", style='List Bullet')
    doc.add_paragraph("Revocations Requested: {{revocations}}", style='List Bullet')
    doc.add_paragraph("Modifications Requested: {{modifications}}", style='List Bullet')
    doc.add_paragraph("New Access Requests: {{new_requests}}", style='List Bullet')

    doc.add_paragraph(
        'All proposed changes were documented and submitted to the Development Team via '
        'Zendesk ticket #{{ticket_number}}. A post-audit follow-up confirmed full implementation '
        '(per Section 7.2 of the Access Control Policy).'
    )

    # Section 4: Authorization
    doc.add_heading('4. Authorization and Delegation', level=2)
    doc.add_paragraph(
        'Per Section 5 of the Access Control Policy, client clinical or program managers are '
        'authorized to assess and determine access requirements for their programs. The attached '
        'access roster was validated and attested by the designated liaison ({{manager_name}}, '
        '{{manager_email}}) the {{clinic_name}} clinic manager, serving as the official '
        'source of truth for this audit cycle.'
    )

    # Section 5: Attachments
    doc.add_heading('5. Attachments', level=2)
    doc.add_paragraph("Attachment A: Reviewed Access Roster – dated {{date_finalized}}", style='List Bullet')

    # Section 6: Compliance Notes
    doc.add_heading('6. Compliance Notes', level=2)
    doc.add_paragraph(
        'This review fulfills the annual access validation requirement defined under Section 7 of the Access Control Policy.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'A follow-up audit will be conducted within 30 days to confirm full implementation.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'This document will be retained for seven years in accordance with our audit log retention schedule.',
        style='List Bullet'
    )

    # Section 7: Signatures
    doc.add_heading('7. Signatures', level=2)

    # Signature table
    sig_table = doc.add_table(rows=2, cols=4)
    sig_table.style = 'Table Grid'

    # Header row
    sig_hdr = sig_table.rows[0].cells
    sig_hdr[0].text = 'Name'
    sig_hdr[1].text = 'Role'
    sig_hdr[2].text = 'Signature'
    sig_hdr[3].text = 'Date'

    # Data row
    sig_data = sig_table.rows[1].cells
    sig_data[0].text = 'Glen Lewis'
    sig_data[1].text = 'VP of Operations / Compliance Lead'
    sig_data[2].text = ''
    sig_data[3].text = ''

    # Save template
    template_path = os.path.expanduser('~/projects/configurations_toolkit/templates/audit_completion_memo.docx')
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    doc.save(template_path)
    print(f"Template created: {template_path}")

if __name__ == '__main__':
    create_template()
